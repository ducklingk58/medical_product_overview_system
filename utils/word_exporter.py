from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from typing import Dict, Any, List

def add_heading_with_numbering(document, text, level=1):
    """번호가 있는 제목을 추가합니다."""
    heading = document.add_heading(text, level=level)
    return heading

def add_bullet_point(paragraph, text, level=0):
    """글머리 기호를 추가합니다."""
    paragraph.add_run(f"• {text}")

def create_structured_overview_document(data: Dict[str, Any], output_path: str):
    """구조화된 개요서 Word 문서를 생성합니다."""
    doc = Document()
    
    # 제목 설정
    title = doc.add_heading('의약품 제품 개요서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3.2.P.1 섹션
    if "3.2.P.1" in data and data["3.2.P.1"]:
        p1_data = data["3.2.P.1"]
        
        # 제품 기본 정보
        doc.add_heading('1. 제품 기본 정보', level=1)
        
        if "product_name" in p1_data:
            doc.add_paragraph(f"제품명: {p1_data['product_name']}")
        
        if "dosage_form" in p1_data:
            doc.add_paragraph(f"제형: {p1_data['dosage_form']}")
            
        if "strength" in p1_data:
            doc.add_paragraph(f"함량: {p1_data['strength']}")
        
        # 외형 정보
        if "appearance" in p1_data:
            doc.add_heading('2. 외형 정보', level=1)
            appearance = p1_data["appearance"]
            if "description" in appearance:
                doc.add_paragraph(f"외형 설명: {appearance['description']}")
            if "identification_mark" in appearance:
                doc.add_paragraph(f"각인 정보: {appearance['identification_mark']}")
        
        # 성분 정보
        if "composition_per_unit" in p1_data:
            doc.add_heading('3. 성분 정보', level=1)
            for i, component in enumerate(p1_data["composition_per_unit"], 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(f"성분명: {component.get('component_name', 'N/A')}")
                p.add_run(f" | 역할: {component.get('role', 'N/A')}")
                p.add_run(f" | 함량: {component.get('amount', 'N/A')}")
                if component.get('standard'):
                    p.add_run(f" | 기준: {component['standard']}")
        
        # 용기 정보
        if "container_closure_system" in p1_data:
            doc.add_heading('4. 용기 정보', level=1)
            container = p1_data["container_closure_system"]
            if "primary" in container:
                primary = container["primary"]
                doc.add_paragraph(f"주용기 재질: {primary.get('material', 'N/A')}")
                doc.add_paragraph(f"주용기 형태: {primary.get('type', 'N/A')}")
            if "secondary" in container and container["secondary"]:
                doc.add_paragraph(f"보조용기: {container['secondary'].get('description', 'N/A')}")
    
    # 3.2.P.2 섹션
    if "3.2.P.2" in data and data["3.2.P.2"]:
        p2_data = data["3.2.P.2"]
        
        # 개발 이력
        if "development_history" in p2_data:
            doc.add_heading('5. 개발 이력', level=1)
            history = p2_data["development_history"]
            
            if "justification_for_formulation" in history:
                doc.add_paragraph(f"제형 선택 근거: {history['justification_for_formulation']}")
            
            if "critical_materials" in history:
                doc.add_heading('5-1. 핵심 원료', level=2)
                for material in history["critical_materials"]:
                    p = doc.add_paragraph()
                    p.add_run(f"• {material.get('name', 'N/A')}: {material.get('impact', 'N/A')}")
            
            if "clinical_batch_info" in history:
                doc.add_heading('5-2. 임상 배치 정보', level=2)
                batch_info = history["clinical_batch_info"]
                if "difference_from_commercial" in batch_info:
                    doc.add_paragraph(f"상업용 배치와의 차이점: {batch_info['difference_from_commercial']}")
        
        # 성능 특성
        if "performance_characteristics" in p2_data:
            doc.add_heading('6. 성능 특성', level=1)
            perf = p2_data["performance_characteristics"]
            
            if "dissolution" in perf:
                doc.add_paragraph(f"용출 특성: {perf['dissolution']}")
            if "disintegration" in perf:
                doc.add_paragraph(f"붕해 특성: {perf['disintegration']}")
            if "bioavailability" in perf:
                doc.add_paragraph(f"생체이용률: {perf['bioavailability']}")
            if "other_properties" in perf:
                doc.add_paragraph("기타 특성:")
                for prop in perf["other_properties"]:
                    doc.add_paragraph(f"• {prop}", style='List Bullet')
        
        # 미생물학적 특성
        if "microbiological_attributes" in p2_data:
            doc.add_heading('7. 미생물학적 특성', level=1)
            micro = p2_data["microbiological_attributes"]
            
            if "preservative" in micro:
                preservative = micro["preservative"]
                doc.add_paragraph(f"보존제: {preservative.get('name', 'N/A')}")
                doc.add_paragraph(f"보존 효과: {preservative.get('efficacy', 'N/A')}")
            
            if "sterility_info" in micro:
                doc.add_paragraph(f"멸균 정보: {micro['sterility_info']}")
    
    # 문서 저장
    doc.save(output_path)

def export_overview_to_word(data: Dict[str, Any], output_path: str):
    """개요서를 Word 문서로 내보냅니다."""
    create_structured_overview_document(data, output_path)
